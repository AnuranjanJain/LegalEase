import re
import difflib
from io import BytesIO
from datetime import datetime
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def generate_redlined_docx(original_text: str, suggested_text: str) -> bytes:
    """
    Compare original_text and suggested_text at the paragraph and word level,
    generating a .docx file where changes are represented as native MS Word
    Track Changes (insertions and deletions).
    """
    doc = Document()
    
    # Enable revision tracking in settings
    try:
        settings = doc.settings.element
        track_revisions = OxmlElement('w:trackRevisions')
        settings.append(track_revisions)
    except Exception:
        pass

    # Split into paragraphs
    orig_paragraphs = original_text.split('\n')
    sugg_paragraphs = suggested_text.split('\n')
    
    matcher = difflib.SequenceMatcher(None, orig_paragraphs, sugg_paragraphs)
    
    revision_id = 0
    now_str = datetime.utcnow().isoformat() + "Z"
    author = "LegalEase AI"

    def add_inserted(p, text):
        nonlocal revision_id
        w_ins = OxmlElement('w:ins')
        w_ins.set(qn('w:id'), str(revision_id))
        revision_id += 1
        w_ins.set(qn('w:author'), author)
        w_ins.set(qn('w:date'), now_str)
        
        w_r = OxmlElement('w:r')
        w_t = OxmlElement('w:t')
        w_t.text = text
        w_r.append(w_t)
        w_ins.append(w_r)
        p._p.append(w_ins)

    def add_deleted(p, text):
        nonlocal revision_id
        w_del = OxmlElement('w:del')
        w_del.set(qn('w:id'), str(revision_id))
        revision_id += 1
        w_del.set(qn('w:author'), author)
        w_del.set(qn('w:date'), now_str)
        
        w_r = OxmlElement('w:r')
        w_delText = OxmlElement('w:delText')
        w_delText.text = text
        w_r.append(w_delText)
        w_del.append(w_r)
        p._p.append(w_del)

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            for idx in range(i1, i2):
                doc.add_paragraph(orig_paragraphs[idx])
        elif tag == 'replace':
            # Sub-diff the content of replacement blocks to keep it fine-grained (word level)
            orig_block = "\n".join(orig_paragraphs[i1:i2])
            sugg_block = "\n".join(sugg_paragraphs[j1:j2])
            
            p = doc.add_paragraph()
            # Split by whitespace but keep the spacing tokens
            orig_tokens = re.split(r'(\s+)', orig_block)
            sugg_tokens = re.split(r'(\s+)', sugg_block)
            
            sub_matcher = difflib.SequenceMatcher(None, orig_tokens, sugg_tokens)
            for sub_tag, si1, si2, sj1, sj2 in sub_matcher.get_opcodes():
                if sub_tag == 'equal':
                    text_chunk = "".join(orig_tokens[si1:si2])
                    if text_chunk:
                        p.add_run(text_chunk)
                elif sub_tag == 'delete':
                    del_text = "".join(orig_tokens[si1:si2])
                    if del_text:
                        add_deleted(p, del_text)
                elif sub_tag == 'insert':
                    ins_text = "".join(sugg_tokens[sj1:sj2])
                    if ins_text:
                        add_inserted(p, ins_text)
                elif sub_tag == 'replace':
                    del_text = "".join(orig_tokens[si1:si2])
                    ins_text = "".join(sugg_tokens[sj1:sj2])
                    if del_text:
                        add_deleted(p, del_text)
                    if ins_text:
                        add_inserted(p, ins_text)
        elif tag == 'delete':
            # Paragraph level deletion
            for idx in range(i1, i2):
                p = doc.add_paragraph()
                add_deleted(p, orig_paragraphs[idx])
        elif tag == 'insert':
            # Paragraph level insertion
            for idx in range(j1, j2):
                p = doc.add_paragraph()
                add_inserted(p, sugg_paragraphs[idx])
                
    out = BytesIO()
    doc.save(out)
    return out.getvalue()
